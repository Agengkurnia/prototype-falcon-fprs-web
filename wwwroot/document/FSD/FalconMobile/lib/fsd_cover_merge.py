"""
Gabungkan cover + revision/approval (2 halaman pertama) dari template FSD standar
dengan konten hasil Pandoc.

Template: templates/FSD_Cover_Template.docx (layout Kalbe Nutritionals)
Logo: templates/logo.png

Digunakan oleh semua build_fsd_*.py di FSD Generator Engine.
"""
import os
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

try:
    from docxcompose.composer import Composer
except ImportError:
    Composer = None

from fsd_paths import COVER_TEMPLATE, LOGO_PATH

BLUE = RGBColor(0x54, 0x8D, 0xD4)
COVER_FONT = 'Arial'
LOGO_WIDTH = Cm(6.51)
LOGO_HEIGHT = Cm(3.93)

# Body child index: simpan 0..20 (cover + revision + page break + TOC SDT)
FRONT_MATTER_KEEP = 21
COVER_TABLE_COUNT = 2

DEFAULT_META = {
    'version': '1.0.0',
    'project': 'IDC System',
    'module': '',
    'module_cover': '',
    'prepared_by': 'Tim ICT',
    'company': 'PT. Sanghiang Perkasa',
    'date': '01/01/2026',
    'brd_no': '2026.SHP-FSD.0040',
    'pid_no': '2026.SHP-PID.0040',
    'revision_version': '1.0.0',
    'revision_date': '1 January 2026',
    'revision_author': 'Tim ICT',
    'revision_desc': 'Initial Document',
}

MONTHS = {
    'jan': '01', 'januari': '01',
    'feb': '02', 'februari': '02',
    'mar': '03', 'maret': '03',
    'apr': '04', 'april': '04',
    'may': '05', 'mei': '05',
    'jun': '06', 'juni': '06',
    'jul': '07', 'juli': '07',
    'aug': '08', 'agustus': '08',
    'sep': '09', 'september': '09',
    'oct': '10', 'oktober': '10',
    'nov': '11', 'november': '11',
    'dec': '12', 'desember': '12',
}


def _normalize_version(v: str) -> str:
    parts = v.strip().split('.')
    while len(parts) < 3:
        parts.append('0')
    return '.'.join(parts[:3])


def _parse_cover_date(text: str) -> str | None:
    """Konversi tanggal MD ke format cover [dd/mm/yyyy]."""
    text = text.strip()
    # 9 April 2026 / 06 Juli 2026
    m = re.match(r'(\d{1,2})\s+(\w+)\s+(\d{4})', text, re.I)
    if m:
        mo = MONTHS.get(m.group(2).lower()[:3]) or MONTHS.get(m.group(2).lower())
        if mo:
            return f"{m.group(1).zfill(2)}/{mo}/{m.group(3)}"
    # Mei 2026 / Apr 2026
    m = re.match(r'(\w+)\s+(\d{4})', text, re.I)
    if m:
        mo = MONTHS.get(m.group(1).lower()[:3]) or MONTHS.get(m.group(1).lower())
        if mo:
            return f"01/{mo}/{m.group(2)}"
    return None


def _parse_revision_rows(md_text: str) -> list[dict]:
    section = re.search(
        r'## Riwayat Revisi\s*\n\n(.+?)(?=\n---|\n## )',
        md_text,
        re.DOTALL,
    )
    if not section:
        return []
    lines = [ln for ln in section.group(1).split('\n') if ln.strip().startswith('|')]
    if len(lines) < 2:
        return []
    rows = []
    for line in lines[2:]:
        parts = [p.strip().strip('*').strip() for p in line.split('|')[1:-1]]
        if len(parts) >= 4 and any(parts) and parts[0] not in ('Versi', '—', '-'):
            rows.append({
                'version': parts[0],
                'date': parts[1],
                'author': parts[2],
                'desc': parts[3],
            })
    return rows


def parse_md_cover_meta(md_text: str, defaults: dict | None = None) -> dict:
    """Ambil metadata cover dari blok header Markdown."""
    meta = {**DEFAULT_META}
    if defaults:
        meta.update(defaults)

    def cell(label: str) -> str | None:
        m = re.search(
            rf'\|\s*\*\*{re.escape(label)}\*\*\s*\|\s*([^|]+?)\s*\|',
            md_text,
        )
        return m.group(1).strip() if m else None

    for label in ('Proyek', 'Project'):
        proj = cell(label)
        if proj:
            meta['project'] = proj
            break

    sistem = re.search(r'^### Sistem:\s*(.+)$', md_text, re.M)
    if sistem and meta['project'] == DEFAULT_META['project']:
        raw = sistem.group(1).strip()
        meta['project'] = raw.split('(')[0].strip() or raw

    v = cell('Versi')
    if v:
        meta['version'] = _normalize_version(v)
        meta['revision_version'] = meta['version']

    t = cell('Tanggal')
    if t:
        meta['revision_date'] = t
        cover_date = _parse_cover_date(t)
        if cover_date:
            meta['date'] = cover_date

    author = cell('Dibuat oleh')
    if author:
        meta['prepared_by'] = author
        meta['revision_author'] = author

    rev_rows = _parse_revision_rows(md_text)
    if rev_rows:
        match = next(
            (r for r in rev_rows if _normalize_version(r['version']) == meta['version']),
            rev_rows[-1],
        )
        meta['revision_version'] = _normalize_version(match['version'])
        meta['revision_date'] = match['date']
        meta['revision_author'] = match['author']
        meta['revision_desc'] = match['desc']

    mod = re.search(r'^## Modul:\s*(.+)$', md_text, re.M)
    if mod:
        full = mod.group(1).strip()
        meta['module'] = full
        for sep in ('—', '–', '-'):
            if sep in full:
                meta['module_cover'] = full.split(sep, 1)[1].strip()
                break
        else:
            # "Item Spec RM (Raw Material Specification)" → singkat sebelum kurung
            short = full.split('(')[0].strip()
            meta['module_cover'] = short or full

    if not meta.get('module_cover'):
        meta['module_cover'] = meta['module'] or meta['project']

    return meta


def strip_md_for_body(md_text: str) -> str:
    """Hanya isi bab (mulai ## 1. Pendahuluan) — tanpa cover/metadata."""
    text = md_text
    text = re.sub(r'^---\s*\n.*?\n---\s*\n+', '', text, count=1, flags=re.DOTALL)
    text = re.sub(r'^# FUNCTIONAL SPECIFICATION.*?\n', '', text, flags=re.DOTALL)
    text = re.sub(r'^## Modul:.*?\n', '', text, flags=re.MULTILINE)
    text = re.sub(r'^### Sistem:.*?\n', '', text, flags=re.MULTILINE)
    text = re.sub(r'^### Versi Dokumen:.*?\n', '', text, flags=re.MULTILINE)
    text = re.sub(r'^---\s*\n', '', text, count=3)
    text = re.sub(
        r'^\| Atribut.*?\n\|[-| ]+\|\n(?:\|.*?\n)+',
        '',
        text,
        count=1,
        flags=re.MULTILINE,
    )
    text = re.sub(
        r'^## Riwayat Revisi\s*\n\n\|.*?(?=\n---\s*\n)',
        '',
        text,
        count=1,
        flags=re.DOTALL,
    )
    text = re.sub(r'^## Daftar Isi\n\n.*?\n---\s*\n\n', '', text, count=1, flags=re.DOTALL)
    m = re.search(r'^## 1\.\s+Pendahuluan', text, re.M)
    if m:
        text = text[m.start():]
    return text.lstrip()


def content_start_index(doc, marker: str = '1. Pendahuluan') -> int:
    """Indeks paragraf pertama isi bab (lewati cover template)."""
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or '').strip()
        if t.startswith(marker):
            return i
    return 25


def _clear_runs(para):
    for r in list(para.runs):
        r._r.getparent().remove(r._r)


def _add_run(para, text, *, size=18, bold=True, italic=False, color=None, font=COVER_FONT):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _set_right_para(para):
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_cell_blue_bracket(cell, text: str):
    cell.text = ''
    para = cell.paragraphs[0]
    _clear_runs(para)
    _add_run(para, f'[{text}]', size=18, bold=True, italic=True, color=BLUE)


def set_cover_logo(doc: Document, logo_path: str = LOGO_PATH):
    if not os.path.exists(logo_path):
        return
    para = doc.paragraphs[0]
    _clear_runs(para)
    run = para.add_run()
    run.add_picture(logo_path, width=LOGO_WIDTH, height=LOGO_HEIGHT)


def update_cover_pages(doc: Document, meta: dict):
    """Update teks cover (paragraf & tabel) sesuai metadata proyek."""
    if len(doc.paragraphs) > 2:
        p = doc.paragraphs[2]
        _clear_runs(p)
        _set_right_para(p)
        _add_run(p, 'Version ', size=18, bold=True, italic=False)
        _add_run(
            p, f"[{meta['version']}]",
            size=18, bold=True, italic=True, color=BLUE,
        )

    if len(doc.paragraphs) > 6:
        p = doc.paragraphs[6]
        _clear_runs(p)
        _set_right_para(p)
        _add_run(p, meta['project'], size=26, bold=True, italic=True, color=BLUE)

    if len(doc.paragraphs) > 7:
        p = doc.paragraphs[7]
        _clear_runs(p)
        _set_right_para(p)
        _add_run(p, 'Modul: ', size=16, bold=True, italic=True, color=BLUE)
        _add_run(
            p, meta.get('module_cover', meta['module']),
            size=16, bold=True, italic=True, color=BLUE,
        )

    if len(doc.paragraphs) > 15:
        p = doc.paragraphs[15]
        _clear_runs(p)
        _set_right_para(p)
        _add_run(
            p, f"[{meta['date']}]",
            size=18, bold=True, italic=True, color=BLUE,
        )

    if doc.tables:
        t0 = doc.tables[0]
        if len(t0.rows) > 1 and len(t0.rows[1].cells) > 1:
            _set_cell_blue_bracket(t0.rows[1].cells[1], meta['prepared_by'])

    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if t1.rows:
            brd = f"BRD No : {meta['brd_no']}"
            pid = f"PID Ref. No : {meta['pid_no']}"
            for ci, cell in enumerate(t1.rows[0].cells):
                cell.text = brd if ci < 4 else pid
        if len(t1.rows) > 3:
            r = t1.rows[3]
            vals = [
                meta['revision_version'],
                meta['revision_date'], meta['revision_date'],
                meta['revision_author'], meta['revision_author'],
                meta['revision_desc'], meta['revision_desc'], meta['revision_desc'],
            ]
            for ci, cell in enumerate(r.cells):
                if ci < len(vals):
                    cell.text = vals[ci]


def trim_template_body(doc: Document, keep: int = FRONT_MATTER_KEEP):
    body = doc.element.body
    for child in list(body)[keep:]:
        body.remove(child)


def merge_cover_and_content(
    content_path: str,
    output_path: str,
    meta: dict,
    template_path: str = COVER_TEMPLATE,
    logo_path: str = LOGO_PATH,
):
    if Composer is None:
        raise RuntimeError('docxcompose belum terinstall. Jalankan: py -m pip install docxcompose')
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Template cover tidak ditemukan: {template_path}')

    shutil.copy2(template_path, output_path)
    master = Document(output_path)

    set_cover_logo(master, logo_path)
    update_cover_pages(master, meta)
    trim_template_body(master)

    composer = Composer(master)
    composer.append(Document(content_path))
    composer.save(output_path)
