"""
build_fsd_web_portal.py
=========================
Pipeline untuk membangun FSD Web Portal Falcon FPRS format DOCX.
Menggunakan kroki.io untuk merender diagram Mermaid menjadi PNG,
lalu Pandoc untuk konversi MD -> DOCX, dan python-docx untuk post-processing.

Run:
    py build_fsd_web_portal.py

Requirements:
    pip install python-docx
    pandoc (https://pandoc.org/installing.html) harus tersedia di PATH
"""
import sys
import io
# Force UTF-8 output to avoid cp1252 UnicodeEncodeError on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import argparse
import base64
import os
import re
import shutil
import subprocess
import urllib.request
import urllib.error
import zlib
import zipfile
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROTOTYPE_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..', '..', '..', '..'))
DOCUMENT_DIR = os.path.join(PROTOTYPE_ROOT, 'Document')
DOCS_WEB_DIR = os.path.join(PROTOTYPE_ROOT, 'docs', 'web')
MD_SRC     = os.path.join(SCRIPT_DIR, 'FSD_Falcon_Web_Portal.md')
MD_TMP     = os.path.join(SCRIPT_DIR, '_tmp_web_portal_processed.md')
MASTER_DEEP = os.path.join(SCRIPT_DIR, '_master_data_deep.md')
DOCX_BASENAME = 'FSD_AKS_MAN_POWER_GT_WEB.docx'
DOCX_LOCAL = os.path.join(SCRIPT_DIR, DOCX_BASENAME)
DOCX_OUT = ''
JOB_CONFIG = None


def build_docx_filename() -> str:
    """Format: YYYYMMDD_HHmmss_FSD_AKS_MAN_POWER_GT_WEB.docx"""
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f'{ts}_FSD_AKS_MAN_POWER_GT_WEB.docx'

# Reference DOCX untuk styling
REF_DOCX_CANDIDATES = [
    os.path.join(SCRIPT_DIR, 'reference.docx'),
    os.path.join(os.path.dirname(SCRIPT_DIR), 'reference.docx'),
]
REF_DOCX = next((p for p in REF_DOCX_CANDIDATES if os.path.exists(p)), None)

SCREENSHOTS = os.path.join(SCRIPT_DIR, 'screenshots')
os.makedirs(SCREENSHOTS, exist_ok=True)

# Output diagram paths
FLOW_MAIN_PNG     = os.path.join(SCREENSHOTS, 'web_portal_flow_main.png')
FLOW_APPROVAL_PNG = os.path.join(SCREENSHOTS, 'web_portal_flow_approval.png')
ERD_PNG           = os.path.join(SCREENSHOTS, 'web_portal_erd.png')
ERD_EXTENDED_PNG  = os.path.join(SCREENSHOTS, 'web_portal_erd_extended.png')

os.makedirs(DOCUMENT_DIR, exist_ok=True)
os.makedirs(DOCS_WEB_DIR, exist_ok=True)

# ──────────────────────────────────────────────────
# STYLE CONSTANTS (sesuai standar FSD IDC System)
# ──────────────────────────────────────────────────
HEADER_BG       = 'D9EAD3'   # light green
BORDER_COLOR    = '000000'   # black
FONT_NAME       = 'Calibri'
FONT_SIZE_BODY  = 11
FONT_SIZE_TABLE = 9


# ──────────────────────────────────────────────────
# STEP 1: Render Mermaid diagrams via Kroki.io
# ──────────────────────────────────────────────────

def render_kroki(mermaid_code: str, output_path: str, label: str) -> bool:
    """Render Mermaid code to PNG via Kroki.io API."""
    try:
        compressed = zlib.compress(mermaid_code.strip().encode('utf-8'), 9)
        b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{b64}'
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        print(f'   [{label}] [OK] {len(data):,} bytes -> {os.path.basename(output_path)}')
        return True
    except Exception as e:
        print(f'   [{label}] [X] Kroki gagal: {e}')
        return False


def step1_render_diagrams(md_content: str):
    """Extract all mermaid blocks and render each to PNG."""
    print('\n[STEP 1] Render diagram Mermaid via Kroki.io...')

    blocks = list(re.finditer(r'```mermaid\s*\n(.*?)```', md_content, flags=re.DOTALL))
    print(f'   Ditemukan {len(blocks)} blok Mermaid.')

    if not blocks:
        print('   [WARN] Tidak ada blok mermaid – diagram dilewati.')
        return

    for i, match in enumerate(blocks):
        code = match.group(1).strip()

        if 'flowchart LR' in code and 'subgraph SL' in code:
            render_kroki(code, FLOW_MAIN_PNG, 'Flow-Main-Swimlane')
        elif 'flowchart LR' in code and 'Creator' in code and 'WaitApproval' in code:
            render_kroki(code, FLOW_APPROVAL_PNG, 'Flow-Approval')
        elif 'erDiagram' in code and 'Tr_Kunjungan' in code:
            render_kroki(code, ERD_EXTENDED_PNG, 'ERD-Extended')
        elif 'erDiagram' in code and 'M_Pelanggan' in code:
            render_kroki(code, ERD_PNG, 'ERD-Main')
        else:
            generic = os.path.join(SCREENSHOTS, f'web_portal_diagram_{i + 1}.png')
            render_kroki(code, generic, f'Diagram-{i + 1}')


# ──────────────────────────────────────────────────
# STEP 2: Preprocess Markdown
# ──────────────────────────────────────────────────

def step0_generate_master_deep():
    print('\n[STEP 0] Generate deep Master Data spec (§7.2)...')
    gen_script = os.path.join(SCRIPT_DIR, 'generate_master_data_deep_spec.py')
    cmd = [sys.executable, gen_script]
    if JOB_CONFIG:
        cmd += ['--job-config', JOB_CONFIG]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        print(result.stderr[:600])
        if not JOB_CONFIG:
            raise RuntimeError('generate_master_data_deep_spec.py failed')
        print('   [WARN] deep spec partial failure — lanjut dengan template')
    else:
        print(result.stdout.strip() or '   [OK] Master data deep spec generated')


def merge_master_data_deep(text: str) -> str:
    deep_path = MASTER_DEEP if os.path.exists(MASTER_DEEP) else os.path.join(SCRIPT_DIR, '_job_build.md')
    if not os.path.exists(deep_path):
        print('   [WARN] deep spec MD tidak ditemukan – lewati merge §7.2')
        return text
    with open(deep_path, 'r', encoding='utf-8') as f:
        deep = f.read()
    deep_body = re.sub(r'^# FUNCTIONAL SPECIFICATION[^\n]*\n+', '', deep, count=1)
    deep_body = re.sub(r'^## Riwayat Revisi[\s\S]*?(?=## 7\.|## \d+\.)', '', deep_body, count=1)
    if '<!-- MASTER_DATA_DEEP:' in text:
        return re.sub(
            r'<!-- MASTER_DATA_DEEP:.*?-->\s*',
            deep_body + '\n',
            text,
            flags=re.DOTALL,
        )
    return text.replace('## 7. Spesifikasi Modul Web Portal', '## 7. Spesifikasi Modul Web Portal\n\n' + deep_body, 1)


def fix_list_formatting(text: str) -> str:
    """Pastikan list bullet/numbered tidak menempel jadi satu paragraf di DOCX."""
    lines = text.split('\n')
    out = []
    list_pat = re.compile(r'^(\*|\-|\d+\.)\s')

    for i, line in enumerate(lines):
        if re.match(r'^\*\s+', line):
            line = re.sub(r'^\*\s+', '- ', line)

        out.append(line)

        if i + 1 >= len(lines):
            continue
        nxt = lines[i + 1]
        if not nxt.strip():
            continue
        if list_pat.match(nxt.strip()) or nxt.lstrip().startswith('- '):
            cur = line.strip()
            if cur and not cur.startswith(('#', '|', '-', '*', '>', '```', '![')):
                if not list_pat.match(cur):
                    out.append('')

    return '\n'.join(out)


def strip_html_comments(text: str) -> str:
    return re.sub(r'<!--.*?-->\s*', '', text, flags=re.DOTALL)


def normalize_front_matter(text: str) -> str:
    """Judul cover & riwayat revisi bukan heading agar tidak masuk Word TOC."""
    lines = []
    for line in text.split('\n'):
        s = line.strip()
        if s.startswith('# FUNCTIONAL SPECIFICATION'):
            lines.append('**FUNCTIONAL SPECIFICATION DOCUMENT (FSD)**')
        elif s.startswith('## Modul:'):
            title = re.sub(r'^##\s*', '', line).replace(' {.unnumbered}', '')
            lines.append(f'**{title.strip()}**')
        elif s.startswith('## Sistem:'):
            title = re.sub(r'^##\s*', '', line).replace(' {.unnumbered}', '')
            lines.append(f'**{title.strip()}**')
        elif s.startswith('## Versi Dokumen:'):
            title = re.sub(r'^##\s*', '', line).replace(' {.unnumbered}', '')
            lines.append(f'**{title.strip()}**')
        elif re.match(r'^## Riwayat Revisi', s):
            lines.append('**Riwayat Revisi**')
        else:
            lines.append(line.replace(' {.unnumbered}', ''))
    return '\n'.join(lines)


def strip_manual_toc(text: str) -> str:
    """Hapus Daftar Isi manual – diganti Word TOC field (standar ItemSpec AKS)."""
    text = re.sub(r'## Daftar Isi.*?(?=\n## \d+\.)', '', text, count=1, flags=re.DOTALL)
    text = re.sub(r'<!-- TOC.*?-->\s*', '', text, flags=re.DOTALL)
    return text


COMPACT_IMAGE_PATTERNS = (
    '_add.png', '_edit.png', '_modal_tambah', '_modal_edit',
    '_validation.png', '_delete_confirm', 'faktur_add', 'canvassing_detail',
)
SMALL_POPUP_PATTERNS = ('_popup', 'motoris_popup', 'audit_popup')


def apply_image_widths(text: str) -> str:
    """Perkecil gambar form/edit/validasi/popup di DOCX."""
    def repl(m):
        alt, path = m.group(1), m.group(2).strip()
        base = path.replace('\\', '/').lower()
        if any(p in base for p in SMALL_POPUP_PATTERNS):
            return f'![{alt}]({path}){{width=45%}}'
        if any(p in base for p in COMPACT_IMAGE_PATTERNS):
            return f'![{alt}]({path}){{width=55%}}'
        return m.group(0)

    return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)(?:\{[^}]*\})?', repl, text)


def patch_document_xml(docx_path: str, mutator) -> bool:
    """Edit word/document.xml in-place via zip."""
    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = zin.namelist()
        parts = {n: zin.read(n) for n in names}
    xml = parts['word/document.xml'].decode('utf-8')
    new_xml = mutator(xml)
    if new_xml == xml:
        return False
    parts['word/document.xml'] = new_xml.encode('utf-8')
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, parts[name])
    return True


def reposition_toc_after_front_matter(docx_path: str) -> bool:
    """Pindahkan Word TOC field ke setelah Riwayat Revisi (seperti ItemSpec AKS)."""
    sdt_re = re.compile(
        r'<w:sdt>\s*<w:sdtPr>.*?docPartGallery w:val="Table of Contents".*?</w:sdt>\s*',
        re.DOTALL,
    )

    def mutator(xml: str) -> str:
        m = sdt_re.search(xml)
        if not m:
            return xml
        sdt_block = m.group(0)
        body = xml[:m.start()] + xml[m.end():]
        anchor = re.search(
            r'<w:p>\s*<w:pPr>\s*<w:pStyle w:val="Heading2"/>.*?<w:t[^>]*>1\. Pendahuluan</w:t>',
            body,
            re.DOTALL,
        )
        if not anchor:
            return xml
        pos = anchor.start()
        return body[:pos] + sdt_block + body[pos:]

    return patch_document_xml(docx_path, mutator)


def update_word_toc(docx_path: str) -> bool:
    """Populate TOC + nomor halaman via Microsoft Word COM (sama seperti dokumen AKS)."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        print('   [WARN] pywin32 tidak terpasang – TOC perlu Update Field manual di Word')
        return False
    path = os.path.abspath(docx_path)
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path, ReadOnly=False)
        if doc.TablesOfContents.Count > 0:
            doc.TablesOfContents(1).Update()
        doc.Fields.Update()
        doc.Save()
        return True
    except Exception as e:
        print(f'   [WARN] Word COM TOC update gagal: {e}')
        return False
    finally:
        if doc is not None:
            doc.Close(SaveChanges=False)
        if word is not None:
            word.Quit()


def enable_update_fields(docx_path: str):
    """Word akan auto-update TOC field (hyperlink ke heading) saat dokumen dibuka."""
    with zipfile.ZipFile(docx_path, 'r') as zin:
        buf = BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml' and b'w:updateFields' not in data:
                    insert = b'<w:updateFields w:val="true"/>'
                    if b'</w:settings>' in data:
                        data = data.replace(b'</w:settings>', insert + b'</w:settings>')
                zout.writestr(item, data)
        with open(docx_path, 'wb') as f:
            f.write(buf.getvalue())


def _rel(path: str) -> str:
    return os.path.relpath(path, SCRIPT_DIR).replace('\\', '/')


def step2_preprocess_markdown():
    print('\n[STEP 2] Pre-processing Markdown...')
    with open(MD_SRC, 'r', encoding='utf-8') as f:
        text = f.read()

    text = merge_master_data_deep(text)
    text = normalize_front_matter(text)
    text = fix_list_formatting(text)
    text = strip_manual_toc(text)
    text = apply_image_widths(text)

    # Render all diagrams first
    step1_render_diagrams(text)

    counter = [0]

    def replace_mermaid(m):
        code = m.group(1).strip()
        counter[0] += 1
        i = counter[0]

        if 'flowchart LR' in code and 'subgraph SL' in code:
            png     = FLOW_MAIN_PNG
            caption = 'Business Flow Diagram Swimlane – Falcon FPRS Web Portal'
        elif 'flowchart LR' in code and 'Creator' in code and 'WaitApproval' in code:
            png     = FLOW_APPROVAL_PNG
            caption = 'Approval Flow – Falcon FPRS Web Portal'
        elif 'erDiagram' in code and 'Tr_Kunjungan' in code:
            png     = ERD_EXTENDED_PNG
            caption = 'ERD Extended – Kunjungan, Stok Motor & Kulakan'
        elif 'erDiagram' in code and 'M_Pelanggan' in code:
            png     = ERD_PNG
            caption = 'ERD – Entity Relationship Diagram Falcon FPRS Web Portal'
        else:
            png     = os.path.join(SCREENSHOTS, f'web_portal_diagram_{i}.png')
            caption = f'Diagram {i}'

        if os.path.exists(png):
            return f'\n![{caption}]({_rel(png)})\n'
        else:
            return f'\n> *[{caption} – diagram tidak dapat di-render]*\n'

    text = re.sub(r'```mermaid\s*\n(.*?)```', replace_mermaid, text, flags=re.DOTALL)
    text = strip_html_comments(text)

    with open(MD_TMP, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f'   [OK] Preprocessed MD -> {os.path.basename(MD_TMP)}')


# ──────────────────────────────────────────────────
# STEP 3: Run Pandoc (MD → DOCX)
# ──────────────────────────────────────────────────

def step3_run_pandoc():
    print('\n[STEP 3] Menjalankan Pandoc (MD -> DOCX)...')
    global DOCX_OUT
    DOCX_OUT = os.path.join(DOCUMENT_DIR, build_docx_filename())

    cmd = [
        'pandoc', MD_TMP,
        '-o', DOCX_OUT,
        '--from=markdown+pipe_tables',
        f'--resource-path={SCRIPT_DIR};{SCREENSHOTS}',
        '--toc',
        '--toc-depth=3',
        '-M', 'toc-title=Table of Contents',
    ]

    if REF_DOCX:
        cmd += [f'--reference-doc={REF_DOCX}']
        print(f'   Menggunakan reference.docx: {REF_DOCX}')
    else:
        print('   [WARN] reference.docx tidak ditemukan – menggunakan style default pandoc.')

    result = subprocess.run(cmd, capture_output=True, text=True, cwd=SCRIPT_DIR)
    if result.returncode != 0:
        print('   Pandoc GAGAL!\n', result.stderr[:800])
        raise RuntimeError('pandoc execution error')
    print(f'   [OK] DOCX berhasil: {os.path.basename(DOCX_OUT)}')


# ──────────────────────────────────────────────────
# STEP 4: Post-process DOCX
# ──────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))


def set_cell_bg(cell, color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def apply_font(run, size_pt: int, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = OxmlElement('w:rFonts')
        rPr.insert(0, el)
    el.set(qn('w:ascii'), FONT_NAME)
    el.set(qn('w:hAnsi'), FONT_NAME)
    el.set(qn('w:cs'), FONT_NAME)
    if bold is not None:
        run.bold = bold


def step4_postprocess_docx():
    print('\n[STEP 4] Post-processing DOCX (Fonts, Table Borders, Image Scale)...')
    doc = Document(DOCX_OUT)

    # Global font style
    try:
        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal'].font.size = Pt(FONT_SIZE_BODY)
    except Exception:
        pass

    # Apply font to all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            apply_font(run, FONT_SIZE_BODY)

    border_spec = {"sz": "8", "val": "single", "color": BORDER_COLOR}
    bdr_all = dict(
        top=border_spec, bottom=border_spec,
        start=border_spec, end=border_spec,
        insideH=border_spec, insideV=border_spec
    )

    # Style all tables
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                set_cell_border(cell, **bdr_all)
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        apply_font(run, FONT_SIZE_TABLE, bold=True if is_header else None)

    # Scale images: default max 12 cm; form/validasi/popup lebih kecil (7 cm)
    MAX_WIDTH_DEFAULT = Cm(12)
    MAX_WIDTH_COMPACT = Cm(7)

    for para in doc.paragraphs:
        para_text = para.text.lower()
        is_compact_ctx = any(k in para_text for k in (
            'modal tambah', 'modal edit', 'form tambah', 'form edit',
            'validasi', 'konfirmasi hapus', 'popup',
        ))
        max_w = MAX_WIDTH_COMPACT if is_compact_ctx else MAX_WIDTH_DEFAULT
        for run in para.runs:
            drawings = run._r.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
            )
            for drawing in drawings:
                extent = drawing.find(
                    '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
                )
                if extent is not None:
                    cx = int(extent.get('cx', 0))
                    if cx > max_w.emu:
                        ratio = max_w.emu / cx
                        cy = int(int(extent.get('cy', 0)) * ratio)
                        extent.set('cx', str(max_w.emu))
                        extent.set('cy', str(cy))

    doc.save(DOCX_OUT)

    if reposition_toc_after_front_matter(DOCX_OUT):
        print('   [OK] TOC dipindah setelah Riwayat Revisi')
    else:
        print('   [WARN] TOC tidak dipindahkan – cek struktur dokumen')

    if update_word_toc(DOCX_OUT):
        print('   [OK] TOC di-update (nomor halaman + hyperlink)')

    enable_update_fields(DOCX_OUT)
    print(f'   [OK] DOCX final tersimpan: {os.path.basename(DOCX_OUT)}')
    print('   [OK] Daftar Isi: Word TOC field (standar ItemSpec AKS)')


def step5_publish_artifacts():
    """Salin DOCX ke folder lokal script dan MD ke docs/web."""
    print('\n[STEP 5] Publish artifacts...')
    shutil.copy2(DOCX_OUT, DOCX_LOCAL)
    latest = os.path.join(DOCUMENT_DIR, 'FSD_AKS_MAN_POWER_GT_WEB_LATEST.docx')
    shutil.copy2(DOCX_OUT, latest)
    print(f'   [OK] Salinan lokal: {os.path.basename(DOCX_LOCAL)}')
    print(f'   [OK] Salinan terbaru: Document/FSD_AKS_MAN_POWER_GT_WEB_LATEST.docx')

    md_dest = os.path.join(DOCS_WEB_DIR, 'FSD_Falcon_Web_Portal.md')
    md_source = MD_TMP if os.path.exists(MD_TMP) else MD_SRC
    shutil.copy2(md_source, md_dest)
    print(f'   [OK] MD disalin ke docs/web/{os.path.basename(md_dest)}')
    print(f'   [OK] DOCX utama: {DOCX_OUT}')


# ──────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--job-config', default=None, help='Path to FSD job JSON from Windows worker')
    args = parser.parse_args()
    JOB_CONFIG = args.job_config

    print('=' * 65)
    print('  BUILD START: FSD Web Portal Falcon FPRS')
    if JOB_CONFIG:
        print(f'  Job config: {JOB_CONFIG}')
    print('=' * 65)

    if not os.path.exists(MD_SRC) and not JOB_CONFIG:
        print(f'\n  ERROR: Source MD tidak ditemukan: {MD_SRC}')
        exit(1)

    step0_generate_master_deep()
    if JOB_CONFIG and os.path.exists(os.path.join(SCRIPT_DIR, '_job_build.md')):
        MD_SRC = os.path.join(SCRIPT_DIR, '_job_build.md')
        print(f'  Using job build MD: {MD_SRC}')

    if not os.path.exists(MD_SRC):
        print(f'\n  ERROR: Source MD tidak ditemukan: {MD_SRC}')
        exit(1)

    step2_preprocess_markdown()
    step3_run_pandoc()
    step4_postprocess_docx()
    step5_publish_artifacts()

    # Cleanup temp file
    if os.path.exists(MD_TMP):
        os.remove(MD_TMP)
        print(f'\n   [OK] Temp file dihapus: {os.path.basename(MD_TMP)}')

    print('\n' + '=' * 65)
    print(f'  SELESAI: {DOCX_OUT}')
    print('=' * 65)
